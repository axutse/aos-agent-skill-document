#!/usr/bin/env python3
"""Build a compact public TAIZHOU brand-enterprise white-paper sample.

The sample uses planning assumptions only and does not contain private data.
"""

from __future__ import annotations

import argparse
import json
import math
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

RED = "7E2027"
DEEP_RED = "5B171C"
BLACK = "1D1D1F"
TEXT = "424245"
GRAY = "86868B"
LIGHT = "E5E5E5"
WARM = "F6F3EF"
WHITE = "FFFFFF"
WAN = "C6B39B"
GE = "2A2927"
UI = "6D7885"


def document_font_family() -> str:
    if Path("/System/Library/Fonts/Hiragino Sans GB.ttc").exists():
        return "Hiragino Sans GB"
    if Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc").exists():
        return "Noto Sans CJK SC"
    if Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf").exists():
        return "Arial Unicode MS"
    return "DejaVu Sans"


DOCX_FONT = document_font_family()


def rgb(hex_value: str) -> tuple[int, int, int]:
    value = hex_value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def find_font(bold: bool = False, size: int = 24) -> ImageFont.FreeTypeFont:
    candidates = [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc" if bold else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, size: int, color: str, bold: bool = False) -> None:
    font = find_font(bold, size)
    x0, y0, x1, y1 = box
    lines = text.split("\n")
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        heights.append(bbox[3] - bbox[1])
    total = sum(heights) + 8 * max(0, len(lines) - 1)
    y = y0 + (y1 - y0 - total) / 2
    for line, height in zip(lines, heights):
        width = draw.textlength(line, font=font)
        draw.text((x0 + (x1 - x0 - width) / 2, y), line, font=font, fill=rgb(color))
        y += height + 8


def diagram_architecture(path: Path) -> None:
    image = Image.new("RGB", (1800, 1050), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.text((85, 50), "TAIZHOU BRAND ENTERPRISE ARCHITECTURE", font=find_font(True, 24), fill=rgb(GRAY))
    draw.text((85, 100), "两家公司建立能力，三个品牌创造市场价值。", font=find_font(True, 42), fill=rgb(BLACK))
    draw.rounded_rectangle((610, 190, 1190, 300), radius=28, fill=rgb(RED))
    draw_center(draw, (630, 200, 1170, 290), "领导与战略决策层", 28, WHITE, True)
    companies = [
        (120, 410, 830, 650, "TAIZHOU新材料科技有限公司", "材料研发｜商品企划｜产品设计\n品牌管理｜渠道销售｜用户数据", DEEP_RED),
        (970, 410, 1680, 650, "TAIZHOU服装有限公司", "样衣版型｜工艺标准｜生产制造\n品质追溯｜仓储包装｜履约交付", BLACK),
    ]
    for x0, y0, x1, y1, title, desc, color in companies:
        draw.rounded_rectangle((x0, y0, x1, y1), radius=28, fill=rgb(WARM), outline=rgb(color), width=4)
        draw_center(draw, (x0 + 30, y0 + 25, x1 - 30, y0 + 105), title, 26, BLACK, True)
        draw_center(draw, (x0 + 35, y0 + 110, x1 - 35, y1 - 25), desc, 20, TEXT)
        draw.line((900, 300, (x0 + x1) / 2, y0), fill=rgb(LIGHT), width=5)
    brands = [
        (100, 790, 580, 985, "万棉尚品", "舒适体感\n规模与复购", WAN),
        (660, 790, 1140, 985, "GEERNA", "材质廓形\n利润与高度", GE),
        (1220, 790, 1700, 985, "UIUP", "年轻身体\n内容与增长", UI),
    ]
    for x0, y0, x1, y1, title, desc, color in brands:
        draw.rounded_rectangle((x0, y0, x1, y1), radius=28, fill=rgb(color), outline=rgb(color), width=2)
        title_color = BLACK if color == WAN else WHITE
        draw_center(draw, (x0 + 20, y0 + 15, x1 - 20, y0 + 85), title, 29, title_color, True)
        draw_center(draw, (x0 + 20, y0 + 85, x1 - 20, y1 - 15), desc, 20, title_color, True)
        draw.line((900, 650, (x0 + x1) / 2, y0), fill=rgb(LIGHT), width=4)
    image.save(path)


def diagram_brand_map(path: Path) -> None:
    image = Image.new("RGB", (1800, 1050), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.text((85, 60), "THREE-BRAND STRATEGY", font=find_font(True, 24), fill=rgb(GRAY))
    draw.text((85, 110), "同一产业底座，三套独立的品牌操作系统。", font=find_font(True, 42), fill=rgb(BLACK))
    left, right, top, bottom = 250, 1650, 260, 880
    draw.line((left, bottom, right, bottom), fill=rgb(BLACK), width=4)
    draw.line((left, bottom, left, top), fill=rgb(BLACK), width=4)
    draw.text((left, bottom + 30), "基础功能价值", font=find_font(False, 20), fill=rgb(GRAY))
    draw.text((right - 170, bottom + 30), "高阶品牌价值", font=find_font(False, 20), fill=rgb(GRAY))
    draw.text((70, bottom - 10), "年轻趋势", font=find_font(False, 20), fill=rgb(GRAY))
    draw.text((70, top - 10), "成熟稳定", font=find_font(False, 20), fill=rgb(GRAY))
    points = [
        ("万棉尚品", 0.25, 0.73, WAN),
        ("GEERNA", 0.77, 0.78, GE),
        ("UIUP", 0.70, 0.25, UI),
    ]
    for label, px, py, color in points:
        x = left + (right - left) * px
        y = bottom - (bottom - top) * py
        draw.ellipse((x - 25, y - 25, x + 25, y + 25), fill=rgb(color))
        draw.text((x + 42, y - 24), label, font=find_font(True, 29), fill=rgb(BLACK))
    image.save(path)


def diagram_product(path: Path) -> None:
    image = Image.new("RGB", (1800, 1050), rgb(WHITE))
    draw = ImageDraw.Draw(image)
    draw.text((85, 55), "PRODUCT & MATERIAL SYSTEM", font=find_font(True, 24), fill=rgb(GRAY))
    draw.text((85, 105), "商品不是从设计开始，而是从用户需求开始。", font=find_font(True, 40), fill=rgb(BLACK))
    steps = ["用户洞察", "品牌定位", "商品企划", "面料与版型", "样衣测试", "生产上市", "数据回流"]
    y = 340
    for index, step in enumerate(steps):
        x0 = 55 + index * 248
        draw.rounded_rectangle((x0, y, x0 + 205, y + 120), radius=22, fill=rgb(WARM), outline=rgb(RED if index == 0 else LIGHT), width=3)
        draw_center(draw, (x0 + 12, y + 12, x0 + 193, y + 108), step, 22, BLACK, True)
        if index < len(steps) - 1:
            draw.line((x0 + 205, y + 60, x0 + 240, y + 60), fill=rgb(LIGHT), width=5)
    categories = [
        ("35%", "打底衫"), ("20%", "针织裤"), ("15%", "卫衣"),
        ("10%", "外套"), ("10%", "保暖内衣"), ("10%", "家居服"),
    ]
    y2 = 650
    for index, (value, name) in enumerate(categories):
        x = 120 + index * 275
        draw.text((x, y2), value, font=find_font(True, 38), fill=rgb(RED if index < 2 else BLACK))
        draw.text((x, y2 + 60), name, font=find_font(True, 21), fill=rgb(TEXT))
    draw.text((90, 930), "PLANNING ASSUMPTION / 企划模拟值", font=find_font(True, 16), fill=rgb(GRAY))
    image.save(path)


def set_cell_border(cell, color: str = LIGHT, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_run(run, size: float, bold: bool = False, color: str = TEXT) -> None:
    run.font.name = DOCX_FONT
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), DOCX_FONT)
    fonts.set(qn("w:hAnsi"), DOCX_FONT)
    fonts.set(qn("w:eastAsia"), DOCX_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*rgb(color))


def paragraph(doc: Document, text: str, size: float = 11, bold: bool = False, color: str = TEXT, after: float = 6, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    add_run(r, size, bold, color)


def title(doc: Document, eyebrow: str, headline: str, subtitle: str | None = None) -> None:
    paragraph(doc, eyebrow.upper(), 8.5, True, GRAY, 7)
    headline_size = 22.5 if len(headline.replace(" ", "")) >= 18 else 27
    paragraph(doc, headline, headline_size, True, BLACK, 9)
    if subtitle:
        paragraph(doc, subtitle, 11, False, TEXT, 12)
    paragraph(doc, "━", 15, True, RED, 8)


def bullet_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("●  ")
        add_run(r, 7, True, RED)
        r = p.add_run(item)
        add_run(r, 10.5, False, TEXT)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_run(cell.paragraphs[0].add_run(header), 9.5, True, BLACK)
        set_cell_border(cell, RED, "10")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            add_run(cells[i].paragraphs[0].add_run(str(value)), 9.2, False, TEXT)
            set_cell_border(cells[i])


def new_page(doc: Document) -> None:
    doc.add_page_break()


def set_page(section, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.4)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.05)
        section.right_margin = Cm(1.95)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.5)


def load_brief(path: Path) -> dict:
    if not path.is_file():
        raise SystemExit(f"Brief not found: {path}")
    with path.open("r", encoding="utf-8") as stream:
        brief = json.load(stream)
    required = ["document_title", "english_title", "planning_horizon", "brands", "data_policy"]
    missing = [key for key in required if not brief.get(key)]
    if missing:
        raise SystemExit(f"Brief is missing required fields: {', '.join(missing)}")
    if len(brief["brands"]) != 3:
        raise SystemExit("The TAIZHOU public case requires exactly three brands")
    return brief


def build(output: Path, brief: dict) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    asset_dir = output.parent / "taizhou_sample_assets"
    asset_dir.mkdir(parents=True, exist_ok=True)
    arch = asset_dir / "architecture.png"
    brands = asset_dir / "brand-map.png"
    product = asset_dir / "product.png"
    diagram_architecture(arch)
    diagram_brand_map(brands)
    diagram_product(product)

    doc = Document()
    set_page(doc.sections[0])
    normal = doc.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal_fonts = normal._element.rPr.rFonts
    normal_fonts.set(qn("w:ascii"), DOCX_FONT)
    normal_fonts.set(qn("w:hAnsi"), DOCX_FONT)
    normal_fonts.set(qn("w:eastAsia"), DOCX_FONT)
    normal.font.size = Pt(10.5)

    # Cover
    paragraph(doc, "TAIZHOU", 11, True, GRAY, 70)
    paragraph(doc, brief["document_title"], 37, True, BLACK, 12)
    paragraph(doc, brief["planning_horizon"].replace("-", "—"), 14, True, RED, 70)
    paragraph(doc, "企业战略、组织治理、多品牌经营、\n商品研发、视觉传播与品牌资产体系", 12, False, TEXT, 14)
    paragraph(doc, brief["english_title"], 9, True, GRAY, 0)

    new_page(doc)
    title(doc, "Core thesis", "TAIZHOU不是第四个消费品牌。", "它是多个消费品牌背后的产业、组织与治理系统。")
    paragraph(doc, "公司承载责任。治理建立秩序。品牌创造价值。数据推动进化。", 17, True, RED, 20)
    bullet_list(doc, [
        "两家公司负责研发、制造、品质与履约。",
        "三个品牌分别承担规模、价值和增长任务。",
        "共享能力提高效率，独立表达创造品牌价值。",
    ])

    new_page(doc)
    title(doc, "Contents", "白皮书目录")
    add_table(doc, ["章节", "核心问题"], [
        ["01 企业与治理", "TAIZHOU是什么，谁决策、谁负责、谁交付"],
        ["02 多品牌战略", "为什么需要万棉尚品、GEERNA与UIUP"],
        ["03 商品与材料", "商品如何开发、分流、测试与复盘"],
        ["04 品牌视觉", "三套品牌如何被用户一眼区分"],
        ["05 传播与经营", "内容、渠道、预算与数据如何增长"],
        ["06 五年方向", "2026—2030如何形成品牌生态"],
    ])

    # Enterprise
    new_page(doc)
    title(doc, "01 / Enterprise and governance", "两家公司建立能力，三个品牌创造市场价值。")
    doc.add_picture(str(arch), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_page(doc)
    title(doc, "Corporate philosophy", "好的服装，不是被制造出来的。")
    paragraph(doc, "它从材料、设计、工艺、制造与人的真实体验中共同形成。", 17, False, TEXT, 20)
    add_table(doc, ["文化主张", "经营翻译"], [
        ["正直", "材料有依据，承诺不夸大，问题不回避"],
        ["进取", "持续优化材料、版型、商品与经营标准"],
        ["协作", "商品、设计、制造、渠道和服务共同负责"],
        ["创新", "让材料、工艺和数据转化为新体验"],
        ["品质", "标准清楚、过程可追溯、交付稳定"],
        ["责任", "尊重环境、伙伴、员工和长期价值"],
    ])

    new_page(doc)
    title(doc, "Governance", "方向、规则与结果，必须分层治理。")
    add_table(doc, ["层级", "主要职责"], [
        ["领导与战略决策", "企业方向、品牌组合、年度目标、预算和重大风险"],
        ["公司治理", "组织、财务、供应链、制造效率、数据和风险"],
        ["品牌治理", "定位、商品边界、视觉资产、价格、渠道和品牌健康"],
        ["品牌经营", "商品、内容、销售、用户、库存与损益结果"],
    ])

    # Brands
    new_page(doc)
    title(doc, "02 / Multi-brand strategy", "同一产业底座，三套独立的品牌操作系统。")
    doc.add_picture(str(brands), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    brand_pages = [
        ("WANMIAN / 万棉尚品", "让舒适成为每天都能感受到的价值。", WAN, ["25—45岁成熟日常女性", "体感舒适生活方式女装", "规模、复购和稳定现金流", "温暖、自然、可靠、真实体感"]),
        ("GEERNA / 哥尔纳", "材料决定触感，廓形决定态度。", GE, ["25—45岁品质都市女性", "极简高阶女装", "品牌高度、客单、利润", "材料、廓形、建筑、长期衣橱"]),
        ("UIUP", "身体向上，风格向前。", UI, ["18—25岁年轻趋势女性", "年轻体感轻户外女装", "新客、内容和增长", "身体、城市、动态、轻机能"]),
    ]
    for brand, statement, color, bullets in brand_pages:
        new_page(doc)
        paragraph(doc, brand, 14, True, color, 35)
        paragraph(doc, statement, 29, True, BLACK, 20)
        bullet_list(doc, bullets)

    new_page(doc)
    title(doc, "Brand firewall", "可以共享能力，但不能模糊品牌边界。")
    add_table(doc, ["共享", "必须独立"], [
        ["材料研发、制造、供应链、品质、数据", "定位、用户、商品、价格、视觉、内容、渠道"],
        ["专业团队和基础技术能力", "核心模特、主场景、详情页母版、包装系统"],
        ["经营数据平台", "品牌用户资产和日常经营目标"],
    ])

    # Product
    new_page(doc)
    title(doc, "03 / Product and material", "商品不是从设计开始，而是从用户需求开始。")
    doc.add_picture(str(product), width=Inches(6.65))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_page(doc)
    title(doc, "Material platforms", "材料不是宣传概念，而是用户可以直接感受到的差异。")
    add_table(doc, ["平台", "代表材料", "产品价值"], [
        ["高阶天然纤维", "牦牛绒、羊毛羊绒", "品牌高度、材质故事、高客单"],
        ["稳定羊毛", "抗起球羊毛、精纺羊毛", "品质、耐穿、长期衣橱"],
        ["体感保暖", "新德绒、空气层、轻暖复合", "轻暖、结构、秋冬规模"],
        ["亲肤肌底", "莫代尔、莱赛尔及混纺", "亲肤、垂坠、跨季复购"],
    ])

    new_page(doc)
    title(doc, "Product evidence", "每一个品牌，都必须有属于自己的代表产品。")
    add_table(doc, ["品牌", "代表产品", "购买理由"], [
        ["万棉尚品", "体感打底、舒适针织裤、空气层卫衣", "亲肤、修饰、稳定、日常复购"],
        ["GEERNA", "材质打底、廓形针织裤、极简外套", "材料、结构、比例、长期价值"],
        ["UIUP", "修身打底、动态针织裤、轻户外套装", "身体比例、弹力、城市与内容传播"],
    ])

    # Visual
    new_page(doc)
    title(doc, "04 / Brand visual systems", "一个企业视觉系统，三套品牌视觉语言。")
    add_table(doc, ["系统", "视觉关键词", "主色方向"], [
        ["TAIZHOU", "材料、科技、秩序、准确", "白、黑、灰、深红"],
        ["万棉尚品", "温暖、自然、生活、体感", "暖白、燕麦、肤色、深咖"],
        ["GEERNA", "材质、空间、廓形、克制", "石色、深棕、黑、银灰"],
        ["UIUP", "身体、城市、动态、轻机能", "冷白、石墨、银灰、少量亮色"],
    ])

    new_page(doc)
    title(doc, "Apple-inspired editorial", "大标题表达判断，留白建立秩序。")
    bullet_list(doc, [
        "一页只表达一个核心观点。",
        "结论先行，结构和证据随后。",
        "核心图占有效页面宽度82%—90%。",
        "减少圆角卡片、中心放射图、渐变和阴影。",
        "品牌先让人感受到，再解释规则，最后证明落地。",
    ])

    # Media and operation
    new_page(doc)
    title(doc, "05 / Media and operation", "从品牌认知，到购买、体验与复购。")
    add_table(doc, ["阶段", "主要内容"], [
        ["认知", "企业实力、品牌态度、场景与人物"],
        ["兴趣", "穿搭、材料故事、动态内容"],
        ["理解", "面料、版型、工艺、功能与价格价值"],
        ["转化", "详情页、达人、直播、私域与客服"],
        ["体验", "包装、履约、穿着、售后"],
        ["复购", "会员、新色、系列延伸、用户分享"],
    ])

    new_page(doc)
    title(doc, "Planning assumptions", "品牌经营必须能够独立核算。")
    add_table(doc, ["品牌", "年度上市款", "核心目标", "价格带（CNY）"], [
        ["万棉尚品", "110", "规模、复购、库存效率", "79—499"],
        ["GEERNA", "70", "客单、毛利、品牌资产", "199—1999"],
        ["UIUP", "60", "新客、内容、测试效率", "99—699"],
    ])
    paragraph(doc, "PLANNING ASSUMPTION / 企划模拟值，正式执行前以年度预算与审批结果为准。", 8.5, True, GRAY, 0)

    # Roadmap
    new_page(doc)
    title(doc, "06 / Five-year direction", "五年不是简单增长，而是品牌能力逐级形成。")
    add_table(doc, ["年度", "阶段", "重点"], [
        ["2026", "基础建设", "定位、治理、视觉、商品母版和市场测试"],
        ["2027", "市场验证", "用户、产品、价格、渠道和品牌经济性"],
        ["2028", "体系完善", "商品、制造、传播、数据和治理标准化"],
        ["2029", "规模复制", "复制成熟产品与有效渠道"],
        ["2030", "品牌生态", "材料、商品、品牌、用户和数据闭环"],
    ])

    new_page(doc)
    paragraph(doc, "TAIZHOU", 11, True, GRAY, 55)
    paragraph(doc, "FROM INDUSTRIAL CAPABILITY\nTO BRAND VALUE", 34, True, BLACK, 28)
    paragraph(doc, "从产业能力出发，\n让商品形成价值，\n让品牌连接用户，\n让数据推动下一轮增长。", 17, False, TEXT, 65)
    paragraph(doc, "WANMIAN  ·  GEERNA  ·  UIUP", 11, True, RED, 8)
    paragraph(doc, "TAIZHOU BRAND ENTERPRISE WHITE PAPER", 8.5, False, GRAY, 0)

    doc.core_properties.title = brief["document_title"]
    doc.core_properties.subject = "Public example for a multi-brand fashion enterprise"
    doc.core_properties.author = "TAIZHOU"
    doc.core_properties.keywords = "TAIZHOU, brand white paper, fashion enterprise, " + ", ".join(brief["brands"])
    doc.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the public TAIZHOU white-paper example.")
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--brief", type=Path, default=Path(__file__).with_name("brief.json"))
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build(args.output, load_brief(args.brief.expanduser().resolve()))
    print(args.output)
