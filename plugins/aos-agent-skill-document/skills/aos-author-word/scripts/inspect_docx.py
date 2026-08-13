#!/usr/bin/env python3
import argparse
import json
from collections import Counter
from pathlib import Path
from zipfile import ZipFile
from lxml import etree
from docx import Document

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def main() -> int:
    parser = argparse.ArgumentParser(description="Inspect DOCX structure, sections, styles, images, and fidelity settings.")
    parser.add_argument("input")
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()

    src = Path(args.input).expanduser().resolve()
    if not src.is_file():
        raise SystemExit(f"Input DOCX not found: {src}")
    document = Document(src)
    styles = Counter(p.style.name for p in document.paragraphs if p.style is not None)

    with ZipFile(src) as zf:
        names = set(zf.namelist())
        media = [name for name in names if name.startswith("word/media/")]
        settings = etree.fromstring(zf.read("word/settings.xml"))
        do_not_compress = settings.find("w:doNotCompressPictures", namespaces=NS) is not None
        dpi = settings.find("w:defaultImageDpi", namespaces=NS)
        default_dpi = dpi.get(f"{{{W_NS}}}val") if dpi is not None else None
        rsid_count = 0
        tracked_changes = 0
        for name in names:
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            root = etree.fromstring(zf.read(name))
            for element in root.iter():
                rsid_count += sum(1 for key in element.attrib if key.startswith(f"{{{W_NS}}}rsid"))
            tracked_changes += len(root.findall(".//w:ins", namespaces=NS))
            tracked_changes += len(root.findall(".//w:del", namespaces=NS))

    properties = document.core_properties

    report = {
        "file": str(src),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "inline_shapes": len(document.inline_shapes),
        "media_files": len(media),
        "metadata": {
            "title": properties.title or "",
            "subject": properties.subject or "",
            "author": properties.author or "",
            "last_modified_by": properties.last_modified_by or "",
            "keywords": properties.keywords or "",
        },
        "has_custom_properties": "docProps/custom.xml" in names,
        "has_comments": "word/comments.xml" in names,
        "tracked_changes": tracked_changes,
        "rsid_attributes": rsid_count,
        "top_paragraph_styles": styles.most_common(15),
        "do_not_compress_pictures": do_not_compress,
        "default_image_dpi": default_dpi,
        "section_details": [
            {
                "index": i + 1,
                "width_inches": round(section.page_width.inches, 3),
                "height_inches": round(section.page_height.inches, 3),
                "orientation": str(section.orientation),
            }
            for i, section in enumerate(document.sections)
        ],
    }
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        for key, value in report.items():
            print(f"{key}: {value}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
