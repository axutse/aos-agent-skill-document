from __future__ import annotations

import json
import shutil
import subprocess
import sys
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
EXAMPLE = ROOT / "examples" / "taizhou-white-paper"
WORD_SCRIPTS = ROOT / "plugins" / "aos-agent-skill-document" / "skills" / "aos-author-word" / "scripts"
PDF_SCRIPTS = ROOT / "plugins" / "aos-agent-skill-document" / "skills" / "aos-process-pdf" / "scripts"


def test_taizhou_example_rebuild_and_metadata_scrub(tmp_path: Path) -> None:
    generated = tmp_path / "taizhou.docx"
    subprocess.run(
        [
            sys.executable,
            str(EXAMPLE / "generate_example.py"),
            "--brief",
            str(EXAMPLE / "brief.json"),
            "--output",
            str(generated),
        ],
        check=True,
    )
    document = Document(generated)
    assert document.core_properties.title == "TAIZHOU品牌企业白皮书"
    assert len(document.paragraphs) >= 100
    assert len(document.tables) == 10

    scrubbed = tmp_path / "taizhou-public.docx"
    subprocess.run(
        [sys.executable, str(WORD_SCRIPTS / "scrub_docx_metadata.py"), str(generated), str(scrubbed), "--author", "TAIZHOU"],
        check=True,
    )
    cleaned = Document(scrubbed)
    assert cleaned.core_properties.author == "TAIZHOU"
    assert not cleaned.core_properties.last_modified_by


def test_pdf_inspection_and_metadata_scrub(tmp_path: Path) -> None:
    source = EXAMPLE / "output" / "TAIZHOU品牌企业白皮书_示例版.pdf"
    inspected = subprocess.run(
        [sys.executable, str(PDF_SCRIPTS / "inspect_pdf.py"), str(source), "--json"],
        check=True,
        capture_output=True,
        text=True,
    )
    report = json.loads(inspected.stdout)
    assert report["pages"] == 20
    assert report["encrypted"] is False

    scrubbed = tmp_path / "public.pdf"
    subprocess.run(
        [
            sys.executable,
            str(PDF_SCRIPTS / "scrub_pdf_metadata.py"),
            str(source),
            str(scrubbed),
            "--title",
            "TAIZHOU品牌企业白皮书",
            "--author",
            "TAIZHOU",
        ],
        check=True,
    )
    reader = PdfReader(scrubbed)
    assert len(reader.pages) == 20
    assert reader.metadata.title == "TAIZHOU品牌企业白皮书"
    assert reader.metadata.author == "TAIZHOU"


@pytest.mark.skipif(shutil.which("pdftoppm") is None, reason="Poppler is not installed")
def test_render_first_pdf_page(tmp_path: Path) -> None:
    source = EXAMPLE / "output" / "TAIZHOU品牌企业白皮书_示例版.pdf"
    subprocess.run(
        [
            sys.executable,
            str(PDF_SCRIPTS / "render_pdf.py"),
            str(source),
            "--output-dir",
            str(tmp_path),
            "--dpi",
            "72",
            "--first-page",
            "1",
            "--last-page",
            "1",
        ],
        check=True,
    )
    assert (tmp_path / "page-001.png").stat().st_size > 1000
